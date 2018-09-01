import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import file_handler

# DONE: CHANGE GENERATION TO MATCH FORMAT

def startup():
    print("1 - Make Cover Sheet")
    print("2 - Create New Style (NON-FUNCTIONAL)")
    print("3 - Edit Style (NON-FUNCTIONAL)")
    print("4 - Display Available Styles")
    userin = input("Anything Else to Quit \n")
    if userin == '1':
        display_style(True)
    elif userin == '2':
        create_style()
    elif userin == '3':
        select_edit()
    elif userin == '4':
        display_style(False)


def display_style(selecting_style=True):
    # DONE: ADD ABILITY TO SELECT STYLE FROM PREDEFINED LIST OF STYLES
    # DONE A: LIST EXISTING FILES IN DIRECTORY
    # DONE B: ALLOW FOR FILE SELECTION
    # DONE: ADD CALL TO SHEET GENERATION
    # On the note of the above task, currently this just runs the default style I made
    styles = []
    for file in os.listdir("Styles"):
        if file.endswith(".style"):
            styles.append(file)
    print("Available Styles Are:")
    for k in range(len(styles)):
        print(str(k+1) + '-', styles[k][:-6])
    if selecting_style:
        style = "Styles/" + select_style(styles)
        if style is not None:
            sheet_information = file_handler.read_in(style)
            generate(sheet_information)
        else:
            display_style()


def select_style(styles):
    try:
        userin = eval(input("Input number of desired style\n"))
        if userin <= len(styles) and not userin == 0:
            return styles[(userin-1)]
        else:
            print('INVALID SELECTION')
            return None
    except NameError:
        print('Invalid Input')
        return None


def create_style():
    # TODO: Add code to allow creation of new styles
    pass


def select_edit():
    # TODO Add ability to edit existing styles
    pass


def input_problems():
    problems = []
    while True:
        temp = input("Enter Problem Titles, !q to quit \n")
        if not temp == "!q":
            problems.append(temp)
        else:
            break
    return problems


def generate(sheet_information):
    document = docx.Document()
    for k in range(len(sheet_information[0])):
        paragraph = document.add_paragraph(sheet_information[0][k])
        center(paragraph)
    if sheet_information[1]:
        set_number = input("Enter Set Number \n")
        set_number = "HW # " + set_number
        paragraph = document.add_paragraph(set_number)
        center(paragraph)
    if sheet_information[3]:
        problems = input_problems()
        if sheet_information[4]:
            problems = [s + '__________' for s in problems]
        for k in range(sheet_information[2]):
            document.add_paragraph()
        for k in range(len(problems)):
            paragraph = document.add_paragraph(problems[k])
            center(paragraph)
    try:
        document.save(sheet_information[6])
    except PermissionError:
        print('ERROR, YOU MUST HAVE THE FILE OPEN!')


"""
def generate(problems, set_number):
    document = docx.Document()
    paragraph = document.add_paragraph("Wesley Turner")
    center(paragraph)
    paragraph = document.add_paragraph("CM 1295")
    center(paragraph)
    paragraph = document.add_paragraph("AC Circuits")
    center(paragraph)
    paragraph = document.add_paragraph("ECE 204-03")
    center(paragraph)
    set_number = "HW # " + str(set_number)
    paragraph = document.add_paragraph(set_number)
    center(paragraph)
    for k in range(10):
        document.add_paragraph()
    for k in range(len(problems)):
        problem_number = str(problems[k]) + '__________'
        paragraph = document.add_paragraph(problem_number)
        center(paragraph)
    try:
        document.save('Cover_Page.docx')
    except PermissionError:
        print("\033[1;31;40mEDIT FAILED! YOU LIKELY HAVE THE COVER_PAGE FILE OPEN! \nCLOSE IT AND TRY AGAIN")
"""


def center(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


if __name__ == "__main__":
    startup()
